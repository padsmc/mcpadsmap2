from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.table import (
    WD_TABLE_ALIGNMENT,
    WD_CELL_VERTICAL_ALIGNMENT,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, color):
    """
    Set the background color of a table cell.

    color example:
    '1F4E78'
    '4472C4'
    """

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)

    tcPr.append(shd)

def create_curriculum_doc(data):
    """
    Generate a Curriculum Map DOCX from the curriculum JSON.
    """

    # ----------------------------------------------------
    # Create Document
    # ----------------------------------------------------

    document = Document()

    section = document.sections[0]

    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)

    section.left_margin = Inches(0.30)
    section.right_margin = Inches(0.30)
    section.top_margin = Inches(0.40)
    section.bottom_margin = Inches(0.40)

    # ----------------------------------------------------
    # Title
    # ----------------------------------------------------

    title = document.add_heading(level=1)

    paragraph = title.paragraph_format
    paragraph.space_after = Pt(12)

    run = title.add_run("CURRICULUM MAP")
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.font.bold = True

    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph()

    # ----------------------------------------------------
    # Standards Table
    # ----------------------------------------------------

    standards = document.add_table(rows=2, cols=3)

    standards.style = "Table Grid"
    standards.alignment = WD_TABLE_ALIGNMENT.CENTER

    standards.autofit = False

    widths = [3.6, 3.6, 3.6]

    for row in standards.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = Inches(width)

    # ----------------------------------------------------
    # Header Row
    # ----------------------------------------------------

    header = standards.rows[0].cells

    header[0].text = "CONTENT STANDARD"
    header[1].text = "PERFORMANCE STANDARD"
    header[2].text = "FORMATION STANDARD"

    # ----------------------------------------------------
    # Data Row
    # ----------------------------------------------------

    values = standards.rows[1].cells

    values[0].text = data.get("content_standard", "")
    values[1].text = data.get("performance_standard", "")
    values[2].text = data.get("formation_standard", "")

    for cell in values:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(10)

    # ----------------------------------------------------
    # Format Standards Table
    # ----------------------------------------------------

    for row in standards.rows:

        for cell in row.cells:

            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(3)
                paragraph.paragraph_format.space_after = Pt(3)

            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.05

                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                for run in paragraph.runs:

                    run.font.name = "Calibri"
                    run.font.size = Pt(10)

    document.add_paragraph()

    # ----------------------------------------------------
    # Curriculum Table (Header Only)
    # ----------------------------------------------------

    curriculum = document.add_table(rows=1, cols=8)

    curriculum.style = "Table Grid"
    curriculum.alignment = WD_TABLE_ALIGNMENT.CENTER
    curriculum.autofit = False

    column_widths = [
        1.25,   # Topic
        2.10,   # Competencies
        2.00,   # Enduring Understanding
        1.60,   # Assessment
        1.70,   # Activities
        1.50,   # Resources
        2.00,   # Integration
        1.60    # Core Values
    ]

    for row in curriculum.rows:
        for i, width in enumerate(column_widths):
            row.cells[i].width = Inches(width)

    headers = curriculum.rows[0].cells

    headers[0].text = "CONTENT / TOPIC"
    headers[1].text = "LEARNING COMPETENCIES"
    headers[2].text = "ENDURING UNDERSTANDING\nESSENTIAL QUESTION"
    headers[3].text = "ASSESSMENT"
    headers[4].text = "ACTIVITIES"
    headers[5].text = "RESOURCES"
    headers[6].text = "INTEGRATION"
    headers[7].text = "INSTITUTIONAL CORE VALUES"

    for cell in headers:
        set_cell_background(cell, "1F4E78")
        
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        for paragraph in cell.paragraphs:

            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for run in paragraph.runs:

                run.font.name = "Calibri"
                run.font.size = Pt(10)
                run.font.bold = True

    # ====================================================
    # PART 2 STARTS BELOW THIS LINE
    # (Do not add anything here yet.)
    # ====================================================
        # ----------------------------------------------------
    # Populate Curriculum Table
    # ----------------------------------------------------

    for topic in data.get("topics", []):

        topic_name = topic.get("topic", "")
        resources = topic.get("resources", [])
        integration = topic.get("integration", {})
        rows = topic.get("rows", [])

        start_row = len(curriculum.rows)

        first_row = True

        for competency in rows:

            cells = curriculum.add_row().cells
            row = curriculum.rows[-1]
            row.height = Pt(42)

            # -----------------------------
            # CONTENT / TOPIC
            # -----------------------------

            if first_row:
                cells[0].text = topic_name.upper()
            else:
                cells[0].text = ""

            # -----------------------------
            # LEARNING COMPETENCIES
            # -----------------------------

            competency_text = (
                f"Level: {competency.get('level', '')}\n\n"
                f"{competency.get('competency', '')}"
            )

            cells[1].text = competency_text

            # -----------------------------
            # ENDURING UNDERSTANDING
            # -----------------------------

            eu = competency.get("enduring_understanding", "")
            eq = competency.get("essential_question", "")

            cells[2].text = (
                f"Enduring Understanding:\n{eu}\n\n"
                f"Essential Question:\n{eq}"
            )

            # -----------------------------
            # ASSESSMENT
            # -----------------------------

            assessments = competency.get("assessment", [])

            if assessments:
                cells[3].text = "\n".join(
                    f"• {item}" for item in assessments
                )

            # -----------------------------
            # ACTIVITIES
            # -----------------------------

            activities = competency.get("activities", [])

            if activities:
                cells[4].text = "\n".join(
                    f"• {item}" for item in activities
                )

            # -----------------------------
            # RESOURCES
            # -----------------------------

            if first_row:

                cells[5].text = "\n".join(
                    f"• {item}" for item in resources
                )

            else:

                cells[5].text = ""

            # -----------------------------
            # INTEGRATION
            # -----------------------------

            if first_row:

                cells[6].text = (
                    f"21st Century Skills\n"
                    f"{integration.get('skill_21st','')}\n\n"
                    f"Vertical Alignment\n"
                    f"{integration.get('vertical_alignment','')}\n\n"
                    f"Horizontal Alignment\n"
                    f"{integration.get('horizontal_alignment','')}"
                )

            else:

                cells[6].text = ""

            # -----------------------------
            # CORE VALUES
            # -----------------------------

            core = competency.get("core_value", {})

            cells[7].text = (
                f"{core.get('value','')}\n\n"
                f"{core.get('reason','')}"
            )

            # -----------------------------
            # Format this row
            # -----------------------------

            for cell in cells:

                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

                for paragraph in cell.paragraphs:

                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    for run in paragraph.runs:

                        run.font.name = "Calibri"
                        run.font.size = Pt(10)

            first_row = False

        # ----------------------------------------
        # Merge Topic, Resources, and Integration
        # ----------------------------------------

        end_row = len(curriculum.rows) - 1

        if end_row > start_row:

            # Topic
            curriculum.cell(start_row, 0).merge(
                curriculum.cell(end_row, 0)
            )

            # Resources
            curriculum.cell(start_row, 5).merge(
                curriculum.cell(end_row, 5)
            )

            # Integration
            curriculum.cell(start_row, 6).merge(
                curriculum.cell(end_row, 6)
            )
            # Align merged cells to the top
            curriculum.cell(start_row, 0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            curriculum.cell(start_row, 5).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            curriculum.cell(start_row, 6).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    # ====================================================
    # PART 3 STARTS BELOW THIS LINE
    # ====================================================
        # ----------------------------------------------------
    # Final Formatting
    # ----------------------------------------------------

    for table in document.tables:

        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for row in table.rows:

            for cell in row.cells:

                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

                for paragraph in cell.paragraphs:

                    if paragraph.alignment is None:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    for run in paragraph.runs:

                        run.font.name = "Calibri"
                        run.font.size = Pt(10)

    # ----------------------------------------------------
    # Save Document
    # ----------------------------------------------------

    output = BytesIO()

    document.save(output)

    output.seek(0)

    return output